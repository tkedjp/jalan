import tkinter
from tkinter import ttk
import datetime
from datetime import timedelta

from time import sleep

import requests
from bs4 import BeautifulSoup
import pandas as pd

import math

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ウィンドウの作成
root = tkinter.Tk()
root.title('じゃらん')
root.geometry('380x270')
root.resizable(0, 0)

#日付
date = tkinter.Label(text='チェックインする日：')
date.grid(row=1, column=1, padx=5, pady=5)
date_box = ttk.Combobox(values=[datetime.date.today() + datetime.timedelta(days=i+1) for i in range(180)])
date_box.grid(row=1, column=2, padx=5, pady=5)

#泊数
stay_count_label = tkinter.Label(text='泊数：')
stay_count_label.grid(row=2, column=1, padx=5, pady=5)
stay_count_box = ttk.Combobox(values=[i+1 for i in range(9)])
stay_count_box.grid(row=2, column=2, padx=5, pady=5)

#室数
room_count_label = tkinter.Label(text='室数：')
room_count_label.grid(row=3, column=1, padx=5, pady=5)
room_count_box = ttk.Combobox(values=[i+1 for i in range(10)])
room_count_box.grid(row=3, column=2, padx=5, pady=5)

#人数
adult_num_label = tkinter.Label(text='人数：')
adult_num_label.grid(row=4, column=1, padx=5, pady=5)
adult_num_box = ttk.Combobox(values=[i+1 for i in range(9)])
adult_num_box.grid(row=4, column=2, padx=5, pady=5)

def save():
    global total_number, processed
    processed = 0
    # max_page_index = 0
    date_box_value = date_box.get()
    month = date_box_value[5:7]
    day = date_box_value[8:10]
    stay_count = stay_count_box.get()
    room_count = room_count_box.get()
    adult_num = adult_num_box.get()

    hotel_list = []
    base_url = 'https://www.jalan.net/040000/LRG_040200/SML_040202/?screenId=UWW1402&distCd=01&listId=0&activeSort=0&mvTabFlg=1&stayYear=2023&stayMonth=' + month + '&stayDay=' + day + '&stayCount=' + stay_count + '&roomCount=' + room_count +'&adultNum=' + adult_num +'&yadHb=1&roomCrack=200000&kenCd=040000&lrgCd=040200&smlCd=040202&vosFlg=6&idx={}'

    user_agent = 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/109.0.0.0 Safari/537.36'
    header = {
        'User-Agent': user_agent
    }

    sleep(3)

    r = requests.get(base_url, timeout=7.5, headers=header)
    if r.status_code >= 400:
        print(F'{base_url}は無効です')

    soup = BeautifulSoup(r.content, 'lxml')

    total_number = soup.select_one('span.jlnpc-listInformation--count').text
    max_page_index = int(total_number) / 30
    max_page_index = round(max_page_index)
    max_page_index = math.floor(max_page_index)

    for i in range(max_page_index):
        url = base_url.format(30*i)

        sleep(3)

        page_r = requests.get(url, timeout=7.5, headers=header)
        if page_r.status_code >= 400:
            print(F'{url}は無効です')
            continue

        page_soup = BeautifulSoup(page_r.content, 'lxml')

        #最安料金と1人あたりの料金
        table_soup = page_soup.select('div.p-yadoCassette__body.p-searchResultItem__body')
        for table in table_soup:
            hotel = table.select_one('a > div > div > div.p-searchResultItem__summaryInner > div.p-searchResultItem__summaryLeft > h2').text
            room_price = table.select_one('a > div > div > div.p-searchResultItem__summaryInner > div.p-searchResultItem__summaryRight > dl > dd > span.p-searchResultItem__lowestPriceValue').text         
            per_price_tag = table.select_one('a > div > div > div.p-searchResultItem__summaryInner > div.p-searchResultItem__summaryRight > dl > dd > span.p-searchResultItem__lowestUnitPrice')
            if per_price_tag is None:
                per_price = None

            else:
                per_price = per_price_tag.text
            
            page_urls = table.select('a.jlnpc-yadoCassette__link')

            for i, page_url in enumerate(page_urls):  
                page_url = 'https://www.jalan.net' + page_url.get('href')
                if 'javascript' in page_url:
                    page_urls = None
                
                else:
                    sleep(3)

                    hotel_page_r = requests.get(page_url, timeout=7.5, headers=header)
                    if hotel_page_r.status_code >= 400:
                        print(F'{page_url}は無効です')
                        continue
                
                    #ホテル名
                    hotel_page_soup = BeautifulSoup(hotel_page_r.content, 'lxml')
        
                    #住所
                    address_tags = hotel_page_soup.select_one('#jlnpc-main-contets-area > div.shisetsu-accesspartking_body_wrap > table tr:nth-child(1) > td')
                    if address_tags is None:
                        address = None

                    else:
                        address = address_tags.text
                        address = address.replace('大きな地図をみる', '')
                        address = address.strip()
                
                    #駐車場
                    parking_tags = hotel_page_soup.select_one('#jlnpc-main-contets-area > div.shisetsu-accesspartking_body_wrap > table tr:nth-child(3) > td')  
                    if parking_tags is None:
                        parking = None
                    
                    else:
                        parking = parking_tags.text
                        parking = parking.replace('\n','')
                        parking = parking.strip()

                    #タイプ別の室数
                    room_tag = hotel_page_soup.select_one('.shisetsu-roomsetsubi_body')
                    if room_tag is None:
                        single = None
                        double = None
                        twin = None
                        sweet = None
                        total = None

                    else:
                        tags = room_tag.text

                        if '総部屋数' not in tags:
                            single = room_tag.select_one('tr:nth-child(2) > td > div > table tr:nth-child(2) > td:first-child').text
                            double = room_tag.select_one('tr:nth-child(2) > td > div > table tr:nth-child(2) > td:nth-child(2)').text
                            twin = room_tag.select_one('tr:nth-child(2) > td > div > table tr:nth-child(2) > td:nth-child(3)').text
                            sweet = room_tag.select_one('tr:nth-child(2) > td > div > table tr:nth-child(2) > td:last-child').text
                            total = None

                        elif 'シングル' not in tags:
                            single = None
                            double = None
                            twin = None
                            sweet = None
                            total = room_tag.select_one('tr:nth-child(1) > td > div > table tr:nth-child(2) > td:nth-child(5)').text
                            total = total.strip()

                        else:
                            single = room_tag.select_one('tr:nth-child(3) > td > div > table tr:nth-child(2) > td:first-child').text
                            double = room_tag.select_one('tr:nth-child(3) > td > div > table tr:nth-child(2) > td:nth-child(2)').text
                            twin = room_tag.select_one('tr:nth-child(3) > td > div > table tr:nth-child(2) > td:nth-child(3)').text
                            sweet = room_tag.select_one('tr:nth-child(3) > td > div > table tr:nth-child(2) > td:last-child').text
                            total = room_tag.select_one('tr:nth-child(1) > td > div > table tr:nth-child(2) > td:nth-child(5)').text
                            total = total.strip()

                    hotel_list.append({
                        'ホテル名': hotel,
                        '詳細ページ': page_url,
                        '住所': address,
                        'シングル': single,
                        'ダブル': double,
                        'ツイン': twin,
                        'スイート': sweet,
                        '総部屋数': total,
                        '室料': room_price,
                        '1人あたり': per_price,
                        '駐車場': parking
                    })
                    # print(hotel_list[-1])
                processed += 1
                progress_label.config(text=f'残りは{processed}/{total_number}です')
                progress_label.update()

    #csv出力
    df = pd.DataFrame(hotel_list)
    df.to_csv('list.csv', index=False, encoding='utf-8-sig')

    def add_hyperlink(paragraph, url, text):

        """
        パラグラフオブジェクトの中にハイパーリンクを配置する関数です。

        :param paragraph: ハイパーリンクを追加する段落。
        :param url: 必要な url を含む文字列
        :param text: urlに対応するテキストを表示します。
        :param return: ハイパーリンクオブジェクト
        """

        # これは document.xml.rels ファイルへのアクセスを取得し、新しいリレーションIDの値を取得します。
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # w:hyperlinkタグを作成し、必要な値を追加する
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # w:r要素を作成する
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # 新しいw:rPr要素を作成する
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # すべてのxml要素を結合し、必要なテキストをw:r要素に追加する
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink

    # wordファイルを新規作成
    doc = docx.Document()

    #ファイルの取得
    df = pd.read_csv('list.csv')  

    for i, r in df.iterrows():
        if i == '':
            break

        #表題
        doc.add_heading(r['ホテル名'], 0)
        doc.add_paragraph('住所：'+ r['住所'])

        #部屋タイプ内訳をテーブルにする
        table = doc.add_table(rows=1, cols=5)

        header_cells = table.rows[0].cells
        header_cells[0].text = 'シングル'
        header_cells[1].text = 'ダブル'
        header_cells[2].text = 'ツイン'
        header_cells[3].text = 'スイート'
        header_cells[4].text = '総部屋数'

        row_cells = table.add_row().cells
        row_cells[0].text = (str(r['シングル']))
        row_cells[1].text = (str(r['ダブル']))
        row_cells[2].text = (str(r['ツイン']))
        row_cells[3].text = (str(r['スイート']))
        row_cells[4].text = (str(r['総部屋数']))

        doc.add_paragraph('')
        doc.add_paragraph('室料：'+ str(r['室料']))
        doc.add_paragraph('1人あたり：'+ str(r['1人あたり']))
        doc.add_paragraph('駐車場：'+ r['駐車場'].strip())
        p = doc.add_paragraph('詳細ページ：'+ r['詳細ページ'])
        add_hyperlink(p, r['詳細ページ'], r['詳細ページ'])

        # 改ページ
        doc.add_page_break()

    # すべての行を左揃えにする
    for p in doc.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    # ファイルを保存
    doc.save('hotel_list.docx')
    
    progress_label.config(text='取得完了')
    progress_label.update()

#処理中
progress = tkinter.Label(text='処理状況：')
progress.grid(row=5, column=1, padx=5, pady=5)
progress_label = tkinter.Label(text='')
progress_label.grid(row=5, column=2, padx=5, pady=5)

#実行
save_button = tkinter.Button(text='取得', command=save)
save_button.grid(row=6, column=2, padx=5, pady=20, ipadx=4, ipady=4)

# ウィンドウのループ処理
root.mainloop()